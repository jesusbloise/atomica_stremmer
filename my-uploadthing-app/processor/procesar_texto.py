import os
import sys
import requests
import tempfile
import psycopg2
import docx  # python-docx
from datetime import datetime
from pathlib import Path
from typing import Optional, List

# ----------------- Consola UTF-8 (arregla 'charmap' en Windows) -----------------
try:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    os.environ.setdefault("PYTHONIOENCODING", "utf-8")
except Exception:
    pass

# ----------------- Config DB (nueva) -----------------
DB_CONFIG = {
    "dbname": os.getenv("PGDATABASE", "falabella_stremmer"),
    "user": os.getenv("PGUSER", "postgres"),
    "password": os.getenv("PGPASSWORD", "atomica"),
    "host": os.getenv("PGHOST", "localhost"),
    "port": os.getenv("PGPORT", "5432"),
}

# ----------------- Utilidades tipo / paths -----------------
def _ext_from(name_or_url: str) -> str:
    base = name_or_url.split("?")[0].split("#")[0]
    return Path(base).suffix.lower()

def _guess_kind(file_name: str, content_type: Optional[str] = None) -> str:
    ct = (content_type or "").lower()
    ext = _ext_from(file_name)

    if ct.startswith("application/pdf") or ext == ".pdf":
        return "pdf"
    if "wordprocessingml.document" in ct or ext == ".docx":
        return "docx"
    if ct.startswith("text/") or ext in {".txt", ".md", ".csv", ".log", ".srt", ".vtt"}:
        return "txt"
    return "desconocido"

def _download_to_temp(url: str, suffix: str) -> str:
    with requests.get(url, stream=True, timeout=60) as r:
        r.raise_for_status()
        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        with os.fdopen(fd, "wb") as f:
            for chunk in r.iter_content(chunk_size=1024 * 64):
                if chunk:
                    f.write(chunk)
    return tmp_path

# ----------------- Extractores -----------------
def _extract_docx(path_docx: str) -> List[str]:
    d = docx.Document(path_docx)
    return [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]

def _extract_pdf(path_pdf: str) -> List[str]:
    text = []
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(path_pdf) as pdf:
            for page in pdf.pages:
                t = page.extract_text() or ""
                t = t.strip()
                if t:
                    text.append(t)
        return text
    except Exception as e:
        print(" Aviso PDF: usando fallback pypdf ->", e)
        try:
            from pypdf import PdfReader  # type: ignore
            reader = PdfReader(path_pdf)
            for p in reader.pages:
                t = p.extract_text() or ""
                t = t.strip()
                if t:
                    text.append(t)
            return text
        except Exception as e2:
            print(" Error PDF (pypdf):", e2)
            return []

def _read_text_with_fallback(path_txt: str) -> str:
    try:
        import chardet  # type: ignore
        with open(path_txt, "rb") as fb:
            raw = fb.read()
        enc = chardet.detect(raw).get("encoding") or "utf-8"
        return raw.decode(enc, errors="replace")
    except Exception:
        for enc in ("utf-8", "latin-1", "cp1252"):
            try:
                with open(path_txt, "r", encoding=enc, errors="replace") as f:
                    return f.read()
            except Exception:
                continue
    return ""

def _extract_txt_like(path_txt: str) -> List[str]:
    content = _read_text_with_fallback(path_txt)
    content = content.replace("\r\n", "\n")
    blocks = [b.strip() for b in content.split("\n\n") if b.strip()]
    return blocks if blocks else ([content.strip()] if content.strip() else [])

# ----------------- Cargar texto multi-formato -----------------
def cargar_texto(
    file_path_or_url: str,
    file_name_hint: Optional[str] = None,
    content_type_hint: Optional[str] = None,
) -> List[str]:
    name_for_type = file_name_hint or file_path_or_url
    kind = _guess_kind(name_for_type, content_type_hint)

    suffix = {
        "docx": ".docx",
        "pdf": ".pdf",
        "txt": ".txt",
        "desconocido": ".bin",
    }.get(kind, ".bin")

    is_url = file_path_or_url.startswith(("http://", "https://"))
    tmp_path = None
    local_path = None

    try:
        if is_url:
            print(f" Descargando archivo desde: {file_path_or_url}")
            tmp_path = _download_to_temp(file_path_or_url, suffix=suffix)
            local_path = tmp_path
        else:
            if not os.path.exists(file_path_or_url):
                raise FileNotFoundError(f"Archivo no encontrado: {file_path_or_url}")
            local_path = file_path_or_url

        if kind == "docx":
            return _extract_docx(local_path)
        elif kind == "pdf":
            return _extract_pdf(local_path)
        elif kind == "txt":
            return _extract_txt_like(local_path)
        else:
            print(" Aviso: tipo no soportado para extracción de texto.")
            return []
    finally:
        if is_url and tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass

# ----------------- Métricas -----------------
def contar_palabras(texto: str):
    palabras = texto.split()
    num_lineas = (texto.count("\n") + 1) if texto else 0
    num_frases = sum(texto.count(x) for x in (".", "!", "?"))
    return {
        "num_lineas": num_lineas,
        "num_palabras": len(palabras),
        "num_frases": num_frases,
    }

# ----------------- Main -----------------
def main(upload_id: str):
    conn = None
    try:
        print(" Conectando a la base de datos...")
        print(f"  -> {DB_CONFIG['dbname']}@{DB_CONFIG['host']}:{DB_CONFIG['port']}")
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        content_type = None
        try:
            cur.execute(
                "SELECT file_path, file_name, tipo, NULL::text as content_type FROM uploads WHERE id = %s",
                (upload_id,),
            )
        except Exception:
            cur.execute(
                "SELECT file_path, file_name, tipo FROM uploads WHERE id = %s",
                (upload_id,),
            )

        row = cur.fetchone()
        if not row:
            print(f" No se encontró el documento con ID {upload_id}")
            return

        if len(row) == 4:
            file_path, file_name, tipo, content_type = row
        else:
            file_path, file_name, tipo = row
            content_type = None

        print(f" Procesando archivo: {file_name}")

        parrafos = cargar_texto(
            file_path_or_url=file_path,
            file_name_hint=file_name,
            content_type_hint=content_type,
        )
        texto_extraido = "\n".join(parrafos) if parrafos else ""
        resumen = " ".join(parrafos[:2]) if parrafos else ""
        stats = contar_palabras(texto_extraido)

        cur.execute(
            """
            INSERT INTO documentos_texto (
                upload_id, tipo, texto, file_name, texto_extraido, creado_en,
                num_lineas, num_palabras, num_frases, resumen
            )
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """,
            (
                upload_id,
                (tipo or "documento"),
                texto_extraido,
                file_name,
                texto_extraido,
                datetime.now(),
                stats["num_lineas"],
                stats["num_palabras"],
                stats["num_frases"],
                resumen,
            ),
        )

        conn.commit()
        print(" ✅ Texto procesado y guardado correctamente en 'documentos_texto'")

    except Exception as e:
        print(f" ❌ Error: {e}")
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(" Debes proporcionar el ID del upload como argumento.")
        sys.exit(1)
    main(sys.argv[1])


# # 

# import os
# import sys
# import requests
# import tempfile
# import psycopg2
# import docx  # python-docx
# from datetime import datetime
# from pathlib import Path
# from typing import Optional, List

# # ----------------- Consola UTF‑8 (arregla 'charmap' en Windows) -----------------
# try:
#     if hasattr(sys.stdout, "reconfigure"):
#         sys.stdout.reconfigure(encoding="utf-8", errors="replace")
#     os.environ.setdefault("PYTHONIOENCODING", "utf-8")
# except Exception:
#     pass

# # ----------------- Config DB -----------------
# DB_CONFIG = {
#     "dbname": "Prueba-atomica-gratis",
#     "user": "postgres",
#     "password": "atomica",
#     "host": "prueba-db-1",
#     "port": "5432",
# }

# # ----------------- Utilidades tipo / paths -----------------
# def _ext_from(name_or_url: str) -> str:
#     base = name_or_url.split("?")[0].split("#")[0]
#     return Path(base).suffix.lower()

# def _guess_kind(file_name: str, content_type: Optional[str] = None) -> str:
#     ct = (content_type or "").lower()
#     ext = _ext_from(file_name)

#     if ct.startswith("application/pdf") or ext == ".pdf":
#         return "pdf"
#     if "wordprocessingml.document" in ct or ext == ".docx":
#         return "docx"
#     if ct.startswith("text/") or ext in {".txt", ".md", ".csv", ".log", ".srt", ".vtt"}:
#         return "txt"
#     return "desconocido"

# def _download_to_temp(url: str, suffix: str) -> str:
#     with requests.get(url, stream=True, timeout=60) as r:
#         r.raise_for_status()
#         fd, tmp_path = tempfile.mkstemp(suffix=suffix)
#         with os.fdopen(fd, "wb") as f:
#             for chunk in r.iter_content(chunk_size=1024 * 64):
#                 if chunk:
#                     f.write(chunk)
#     return tmp_path

# # ----------------- Extractores -----------------
# def _extract_docx(path_docx: str) -> List[str]:
#     d = docx.Document(path_docx)
#     return [p.text.strip() for p in d.paragraphs if p.text and p.text.strip()]

# def _extract_pdf(path_pdf: str) -> List[str]:
#     text = []
#     try:
#         import pdfplumber  # type: ignore
#         with pdfplumber.open(path_pdf) as pdf:
#             for page in pdf.pages:
#                 t = page.extract_text() or ""
#                 t = t.strip()
#                 if t:
#                     text.append(t)
#         return text
#     except Exception as e:
#         print(" Aviso PDF: usando fallback pypdf ->", e)
#         try:
#             from pypdf import PdfReader  # type: ignore
#             reader = PdfReader(path_pdf)
#             for p in reader.pages:
#                 t = p.extract_text() or ""
#                 t = t.strip()
#                 if t:
#                     text.append(t)
#             return text
#         except Exception as e2:
#             print(" Error PDF (pypdf):", e2)
#             return []

# def _read_text_with_fallback(path_txt: str) -> str:
#     try:
#         import chardet  # type: ignore
#         with open(path_txt, "rb") as fb:
#             raw = fb.read()
#         enc = chardet.detect(raw).get("encoding") or "utf-8"
#         return raw.decode(enc, errors="replace")
#     except Exception:
#         for enc in ("utf-8", "latin-1", "cp1252"):
#             try:
#                 with open(path_txt, "r", encoding=enc, errors="replace") as f:
#                     return f.read()
#             except Exception:
#                 continue
#     return ""

# def _extract_txt_like(path_txt: str) -> List[str]:
#     content = _read_text_with_fallback(path_txt)
#     content = content.replace("\r\n", "\n")
#     blocks = [b.strip() for b in content.split("\n\n") if b.strip()]
#     return blocks if blocks else ([content.strip()] if content.strip() else [])

# # ----------------- Cargar texto multi-formato -----------------
# def cargar_texto(
#     file_path_or_url: str,
#     file_name_hint: Optional[str] = None,
#     content_type_hint: Optional[str] = None,
# ) -> List[str]:
#     name_for_type = file_name_hint or file_path_or_url
#     kind = _guess_kind(name_for_type, content_type_hint)

#     suffix = {
#         "docx": ".docx",
#         "pdf": ".pdf",
#         "txt": ".txt",
#         "desconocido": ".bin",
#     }.get(kind, ".bin")

#     is_url = file_path_or_url.startswith(("http://", "https://"))
#     tmp_path = None
#     local_path = None

#     try:
#         if is_url:
#             print(f" Descargando archivo desde: {file_path_or_url}")
#             tmp_path = _download_to_temp(file_path_or_url, suffix=suffix)
#             local_path = tmp_path
#         else:
#             if not os.path.exists(file_path_or_url):
#                 raise FileNotFoundError(f"Archivo no encontrado: {file_path_or_url}")
#             local_path = file_path_or_url

#         if kind == "docx":
#             return _extract_docx(local_path)
#         elif kind == "pdf":
#             return _extract_pdf(local_path)
#         elif kind == "txt":
#             return _extract_txt_like(local_path)
#         else:
#             print(" Aviso: tipo no soportado para extracción de texto.")
#             return []
#     finally:
#         if is_url and tmp_path and os.path.exists(tmp_path):
#             try:
#                 os.remove(tmp_path)
#             except Exception:
#                 pass

# # ----------------- Métricas -----------------
# def contar_palabras(texto: str):
#     palabras = texto.split()
#     num_lineas = (texto.count("\n") + 1) if texto else 0
#     num_frases = sum(texto.count(x) for x in (".", "!", "?"))
#     return {
#         "num_lineas": num_lineas,
#         "num_palabras": len(palabras),
#         "num_frases": num_frases,
#     }

# # ----------------- Main -----------------
# def main(upload_id: str):
#     conn = None
#     try:
#         conn = psycopg2.connect(**DB_CONFIG)
#         cur = conn.cursor()

#         content_type = None
#         try:
#             cur.execute(
#                 "SELECT file_path, file_name, tipo, NULL::text as content_type FROM uploads WHERE id = %s",
#                 (upload_id,),
#             )
#         except Exception:
#             cur.execute(
#                 "SELECT file_path, file_name, tipo FROM uploads WHERE id = %s",
#                 (upload_id,),
#             )

#         row = cur.fetchone()
#         if not row:
#             print(f" No se encontró el documento con ID {upload_id}")
#             return

#         if len(row) == 4:
#             file_path, file_name, tipo, content_type = row
#         else:
#             file_path, file_name, tipo = row
#             content_type = None

#         print(f" Procesando archivo: {file_name}")

#         parrafos = cargar_texto(
#             file_path_or_url=file_path,
#             file_name_hint=file_name,
#             content_type_hint=content_type,
#         )
#         texto_extraido = "\n".join(parrafos) if parrafos else ""
#         resumen = " ".join(parrafos[:2]) if parrafos else ""
#         stats = contar_palabras(texto_extraido)

#         cur.execute(
#             """
#             INSERT INTO documentos_texto (
#                 upload_id, tipo, texto, file_name, texto_extraido, creado_en,
#                 num_lineas, num_palabras, num_frases, resumen
#             )
#             VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
#             """,
#             (
#                 upload_id,
#                 (tipo or "documento"),
#                 texto_extraido,
#                 file_name,
#                 texto_extraido,
#                 datetime.now(),
#                 stats["num_lineas"],
#                 stats["num_palabras"],
#                 stats["num_frases"],
#                 resumen,
#             ),
#         )

#         conn.commit()
#         print(" Texto procesado y guardado correctamente en 'documentos_texto'")

#     except Exception as e:
#         print(f" Error: {e}")
#     finally:
#         if conn:
#             conn.close()

# if __name__ == "__main__":
#     if len(sys.argv) < 2:
#         print(" Debes proporcionar el ID del upload como argumento.")
#         sys.exit(1)
#     main(sys.argv[1])

